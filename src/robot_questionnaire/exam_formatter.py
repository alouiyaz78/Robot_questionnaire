from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt


def _now_stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _add_normal(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)


def _add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)


def _md_escape(s: str) -> str:
    return s.replace("\r", "").strip()


def to_exam_markdown(questions: List[Dict[str, Any]]) -> str:
    """
    Rend un markdown style "examen" proche de ton exemple.
    """
    out: List[str] = []
    for i, q in enumerate(questions, start=1):
        pts = q.get("points", 2)
        out.append(f"**Question {i} ({pts} points)**")
        out.append(_md_escape(q["text"]))
        out.append("")
        out.append(f"Options de la question {i} :")
        out.append("")

        qtype = q.get("type")
        if qtype == "true_false":
            out.append("- Vrai")
            out.append("- Faux")

        elif qtype == "mcq_single":
            for opt in q.get("options", []):
                out.append(f"- {opt}")

        elif qtype == "multi_select":
            out.append("(plusieurs réponses possibles)")
            for opt in q.get("options", []):
                out.append(f"- {opt}")

        elif qtype == "matching":
            right = q.get("right", [])
            left = q.get("left", [])
            out.append("Correspondances :")
            out.append("")
            out.append("**Colonne A (contre-mesures)**")
            for idx, r in enumerate(right, start=1):
                out.append(f"{idx}. {r}")
            out.append("")
            out.append("**Colonne B (attaques / éléments)**")
            for idx, l in enumerate(left, start=1):
                out.append(f"{idx}. {l}")

        elif qtype == "open_long":
            out.append(_md_escape(q.get("instruction", "Répondez en 10–15 lignes.")))

        else:
            out.append("- (Type non supporté)")

        out.append("\n---\n")

    return "\n".join(out).strip() + "\n"


def to_answer_markdown(questions: List[Dict[str, Any]]) -> str:
    out: List[str] = []
    for i, q in enumerate(questions, start=1):
        out.append(f"**Question {i}**")
        out.append(_md_escape(q["text"]))
        out.append("")
        out.append("**Réponse :**")

        qtype = q.get("type")
        ans = q.get("answer")

        if qtype in ("true_false", "mcq_single"):
            out.append(str(ans))

        elif qtype == "multi_select":
            # ans = ["A", ...] ou valeurs exactes
            if isinstance(ans, list):
                out.append(", ".join(str(x) for x in ans))
            else:
                out.append(str(ans))

        elif qtype == "matching":
            # ans dict: left -> right
            if isinstance(ans, dict):
                for k, v in ans.items():
                    out.append(f"- {k} → {v}")
            else:
                out.append(str(ans))

        elif qtype == "open_long":
            out.append(_md_escape(q.get("rubric", "(corrigé/rubrique non fourni)")))

        else:
            out.append(str(ans))

        out.append("\n---\n")

    return "\n".join(out).strip() + "\n"


def export_exam_docx(
    filepath: str,
    title: str,
    questions: List[Dict[str, Any]],
    include_answers: bool = False,
    user_answers: Optional[List[Dict[str, Any]]] = None,
) -> None:
    """
    Génère un .docx style examen.
    - include_answers=True: ajoute corrigé
    - user_answers: liste alignée avec questions, pour inclure "ta réponse"
    """
    doc = Document()
    _add_title(doc, title)
    _add_normal(doc, f"Généré le {_now_stamp()}")

    doc.add_paragraph("")

    for i, q in enumerate(questions, start=1):
        pts = q.get("points", 2)
        _add_h2(doc, f"Question {i} ({pts} points)")
        _add_normal(doc, q["text"])

        qtype = q.get("type")

        doc.add_paragraph("")
        _add_normal(doc, f"Options de la question {i} :")

        if qtype == "true_false":
            _add_bullet(doc, "Vrai")
            _add_bullet(doc, "Faux")

        elif qtype == "mcq_single":
            for opt in q.get("options", []):
                _add_bullet(doc, str(opt))

        elif qtype == "multi_select":
            _add_normal(doc, "(plusieurs réponses possibles)")
            for opt in q.get("options", []):
                _add_bullet(doc, str(opt))

        elif qtype == "matching":
            right = q.get("right", [])
            left = q.get("left", [])
            _add_normal(doc, "Colonne A (contre-mesures) :")
            for idx, r in enumerate(right, start=1):
                _add_numbered(doc, f"{r}")
            _add_normal(doc, "Colonne B (attaques / éléments) :")
            for idx, l in enumerate(left, start=1):
                _add_numbered(doc, f"{l}")

        elif qtype == "open_long":
            _add_normal(doc, q.get("instruction", "Répondez en 10–15 lignes."))

        else:
            _add_normal(doc, "(Type non supporté)")

        # Ta réponse
        if user_answers and i - 1 < len(user_answers):
            ua = user_answers[i - 1].get("user_answer")
            if ua:
                doc.add_paragraph("")
                _add_normal(doc, f"Ta réponse : {ua}")

        # Corrigé
        if include_answers:
            doc.add_paragraph("")
            _add_normal(doc, "Corrigé :")
            ans = q.get("answer")

            if qtype == "matching" and isinstance(ans, dict):
                for k, v in ans.items():
                    _add_bullet(doc, f"{k} → {v}")
            elif qtype == "open_long":
                _add_normal(doc, q.get("rubric", "(rubrique non fournie)"))
            else:
                _add_normal(doc, str(ans))

        doc.add_page_break()

    doc.save(filepath)
